#!encoding=utf-8
import os
import time
import requests
from bs4 import BeautifulSoup
# noinspection PyUnresolvedReferences
from urllib.parse import urlparse
# from logger import logger
import re
import urllib
from html2text import HTML2Text
from PIL import Image
import urllib.request
from urllib import request
import base64
import yaml
import random
import datetime
import threading
import shutil
import traceback
from source.utils import *
import argparse
from fake_useragent import UserAgent
from docx import Document



class Spider(object):
    
    def __init__(self,args):
        
        self.ua = UserAgent()
        self.PAGE_NUM = 1
        self.name = '量科'
        self.driver_normal = True

        self.args = args
        self.args.name = self.name
        self.args.industry_name = self.args.industry
        self.args.industry_type = allinfo.get(self.args.industry_name, ['',''])[1]
        self.args.industry_ID = allinfo.get(self.args.industry_name, ['',''])[0]
        self.args.dom_or_int = 'DOM'
        self.args.source_name = '量科'
        self.args.fuzeren = '' 

        try:
            self.args.start = datetime.datetime.strptime(self.args.start, "%Y-%m-%d %H:%M:%S")
            self.args.end = datetime.datetime.strptime(self.args.end, "%Y-%m-%d %H:%M:%S")
        except:
            self.args.start = ''
            self.args.end = ''
        
    def clean_markdown(self, args, idd, md, start_words=[], stop_words=[]):
        '''清洗markdown源数据，下载图片等
        idd：编号
        md：需要清洗的markdown
        start_words：获取到该词在内容中就开始处理
        stop_words：获取到该词在内容中就停止处理
        '''

        os.makedirs("./output/{}/{}/doc/{}/".format(args.day, args.industry, args.name), exist_ok = True)
        os.makedirs("./output/{}/{}/mdFiles/{}/".format(args.day, args.industry, args.name), exist_ok = True)
        os.makedirs("./output/{}/{}/text/{}/".format(args.day, args.industry, args.name), exist_ok = True)
        os.makedirs("./output/{}/{}/thumbFiles/{}/".format(args.day, args.industry, args.name), exist_ok = True) 

        text_list = [text.strip('\n').strip() for text in md.split('\n') if text.strip('\n').strip() != '']
        res = []
        onlytext_res = []
        n = 1
        output_doc = Document() 
        paragraph = output_doc.add_paragraph()
        
        flag = False
        startflag = False
        content_set = set()
        for text in text_list:
            if ('.jpg' in text) or ('.png' in text) or ('.jpeg') in text:                

                try:
                    reg = '([0-9a-zA-Z_/\.-]*?\.jpg)|([0-9a-zA-Z_/\.-]*?\.png)|([0-9a-zA-Z_/\.-]*?\.jpeg)'

                    imgreg = re.compile(reg)
                    if len(re.findall(imgreg, text)) > 0:
                        picture_turple = re.findall(imgreg, text)[0]
                        picture_url = [x.strip() for x in list(picture_turple) if x.strip() != ''][0]
                        picture_url = '	http://www.qtc.com.cn' + picture_url
                    else:
                        picture_url = ''

                    if picture_url == '' or picture_url in content_set:
                        continue
                    content_set.add(picture_url)

                    print('下载第{}张图片:{}'.format(n, picture_url))
                    img_type = 'png'
                    file_path = './tmp/img/{}_pre_{}.{}.{}'.format(args.name,idd,n,img_type)
                    resize_file_path = './tmp/img/{}_{}.{}.{}'.format(args.name,idd,n,img_type)
                    download(file_path, picture_url)
                    resize_image_1000(file_path, resize_file_path)
                    resize_base64_str = image_to_base64(resize_file_path)
                    imagetext = '''![image](data:image/{};base64,{})'''
                    image_base64text = imagetext.format(resize_file_path.split('.')[-1],resize_base64_str)

                    res.append(image_base64text)
                    output_doc.add_picture(file_path)
                    paragraph = output_doc.add_paragraph()

                    # 把第一张存为缩略图
                    if n == 1:
                        image_path = file_path
                        if args.thub_type == 'txt':
                            resize_image_path = './tmp/img/{}_thub_{}.{}.png'.format(args.name,idd,n)

                            result_path = './output/{}/{}/thumbFiles/{}/{}.txt'.format(args.day,
                                                                                        args.industry,
                                                                                        args.name,idd)
                            thumbnail(image_path, resize_image_path, result_path)
                        else:
                            resize_image_path = './output/{}/{}/thumbFiles/{}/{}.png'.format(args.day,
                                                                                            args.industry,
                                                                                            args.name,idd)
                            result_path = resize_image_path = './tmp/img/{}_thub_{}.{}.txt'.format(args.name,idd,n)

                            thumbnail(image_path, resize_image_path, result_path)
                    n += 1

                except Exception as e:
                    print('处理图片出现错误：{}'.format(e))
                    pass
            else:
                try:
                    content = replace_symbol(text)
                    if content in content_set:
                        continue
                    content_set.add(content)


                    if len(start_words) == 0:
                                startflag = True
                    else:
                        for start_word in start_words:
                            if start_word in content:
                                startflag = True

                    if startflag:
                        for stop_word in stop_words:
                            if stop_word in content:
                                flag = True
                                break
                        
                        if flag:
                            break
                        onlytext_res.append(content)
                        if str(content).startswith('## '):
                            res.append('## <font color="#00e4ff">**{}**</font>'.format(content))
                            paragraph.add_run(content).bold = True
                            paragraph.add_run('\n\n')
                        elif str(content).startswith('### '):
                            res.append('### <font color="#00e4ff">**{}**</font>'.format(content))
                            paragraph.add_run(content).bold = True
                            paragraph.add_run('\n\n')
                        else:
                            res.append(content)
                            paragraph.add_run(content+'\n\n')
                except Exception as e:
                    print('处理文字出现错误：{}'.format(e))
                    pass

        res_text = '\n\n'.join(res)
        onlytext_res_text = '\n\n'.join(onlytext_res)
        return res_text, onlytext_res_text, output_doc
    
    #将文章链接保存为md
    def html2res(self, url, date='', i=1):
        time.sleep(random.randint(1, 5))
        response = requests.get(url, headers={'headers':self.ua.random}, verify=False)
        # 解析 HTML
        soup = BeautifulSoup(response.content, 'html.parser')
        # 获取文章正文
        # print(soup)
        title = soup.find('article', {'class': 'article'}).find('h1',{'class':'page-header'}).get_text().strip()
        title = rep_comma(title)
        print('题目:\t', title)

        # 热度、热词, 点赞，评论
        redu, reci, dianzan, pinglun, yuedu = '', '', '', '', ''
        article = soup.find('article', {'class': 'article'}).find('div', {'class': 'content'})
        #去除
        for element in article.find_all('div',{'class':['ad-show', 'tags-share']}):
            element.extract()
        article = article.prettify()
        # print(article)
       
        text_maker = HTML2Text()
        text_maker.body_width = 0
        text_maker.ignore_links = True
        text_maker.ignore_tables = True
        md_text = text_maker.handle(str(article))
        # print(md_text)

        start_words = []
        stop_words = []
        clean_md, clean_text, clean_doc = self.clean_markdown(self.args, i, md_text, start_words=start_words, stop_words=stop_words)

        return clean_md, clean_text, clean_doc, redu, reci, title, date, url, dianzan, pinglun, yuedu
    
    #将文章链接保存为md
    def html2res2(self, url, date='', i=1):
        time.sleep(random.randint(1, 5))
        response = requests.get(url, headers={'headers':self.ua.random}, verify=False)
        # 解析 HTML
        soup = BeautifulSoup(response.content, 'html.parser')
        # 获取文章正文
        # print(soup)
        title = soup.find('div', {'class': 'flash-bg'}).find('div',{'class':'info'}).find('h2').get_text().strip()
        title = rep_comma(title)
        print('题目:\t', title)

        # 热度、热词, 点赞，评论
        redu, reci, dianzan, pinglun, yuedu = '', '', '', '', ''
        article = soup.find('div', {'class': 'flash-bg'}).find('div',{'class':'txt'}).find('p').get_text().strip()
        
        # print(article)
       
        text_maker = HTML2Text()
        text_maker.body_width = 0
        text_maker.ignore_links = True
        text_maker.ignore_tables = True
        md_text = text_maker.handle(str(article))
        # print(md_text)

        start_words = []
        stop_words = []
        clean_md, clean_text, clean_doc = self.clean_markdown(self.args, i, md_text, start_words=start_words, stop_words=stop_words)

        return clean_md, clean_text, clean_doc, redu, reci, title, date, url, dianzan, pinglun, yuedu
    
    def run(self):
        #采集资讯和快讯数据
        url1 = 'http://www.qtc.com.cn/news'    #资讯
        url2 = 'http://www.qtc.com.cn/flash'   #快讯
        i = 0
        idd = 1
        try:
            flag = False
            while i<self.PAGE_NUM:
                if flag:
                    break
                # 随机暂停几秒，避免过快的请求导致过快的被查到
                time.sleep(random.randint(1, 5))
                resp = requests.get(url1, headers={'headers':self.ua.random}, verify=False)
                # print('resp', resp.text)
                soup = BeautifulSoup(resp.content, 'html.parser')
                all_news_box = soup.find('div',{'class':'news-list'}).find_all('li',{'class':'item'})
                # print(all_news_box)
                for new_info in all_news_box:
                    detail_url = ''
                    try:
                        detail_url = new_info.find('h3',{'class','title clearfix'}).find('a')['href']
                        detail_url = 'http://www.qtc.com.cn/' + detail_url

                        # 打开资讯网页判断时间是否符合需求
                        resp_new = requests.get(detail_url, headers={'headers':self.ua.random}, verify=False)
                        soup_new = BeautifulSoup(resp_new.content, 'html.parser')

                        date_ori = soup_new.find('article', {'class': 'article'}).find('div', {'class': 'page-bar'}).find('span',{'class':'time'}).get_text().strip()
                        # print(date_ori)
                        date = date_ori + ':00'
                        print(date)
                        date = datetime.datetime.strptime(str(date).split('.')[0], "%Y-%m-%d %H:%M:%S")

                        if self.args.start <= date <= self.args.end:
                            clean_md, clean_text, clean_doc, redu, reci, title, date, url, dianzan, \
                                pinglun, yuedu = self.html2res(detail_url, date, str(idd))
                            if self.args.keyword == '' or self.args.keyword in clean_text:
                                save_file(self.args, idd, clean_md, clean_text, clean_doc, redu, reci, title, date, url, dianzan, pinglun, yuedu)
                                idd += 1
                            else:
                                print('该文章不包含关键词：{}！跳过！'.format(self.args.keyword))
                        elif date > self.args.end:
                            print('当前资讯比查找范围最新日期还要新，继续查找')
                            continue
                        else:
                            print('当前资讯过早，结束查找！')
                            flag = True
                            break
                    except Exception as e:
                        print(e)
                        print(detail_url + '  download failure !!!')
                        continue
                print(f"第{i}页爬取成功\n")
                i += 1
        except Exception as e:
            print('整体资讯网页爬取异常：%s' % e)
            print(traceback.format_exc())
            pass

        i = 0
        try:
            flag = False
            while i<self.PAGE_NUM:
                if flag:
                    break
                # 随机暂停几秒，避免过快的请求导致过快的被查到
                time.sleep(random.randint(1, 5))
                resp = requests.get(url2, headers={'headers':self.ua.random}, verify=False)
                # print('resp', resp.text)
                soup = BeautifulSoup(resp.content, 'html.parser')
                all_news_box = soup.find('div',{'class':'flash-list'}).find_all('div',{'class':'item clearfix'})
                # print(all_news_box)
                for new_info in all_news_box:
                    detail_url = ''
                    try:
                        detail_url = new_info.find('div',{'class':['txt hot-0','txt hot-1']}).find('a')['href']
                        detail_url = 'http://www.qtc.com.cn/' + detail_url

                        # 打开资讯网页判断时间是否符合需求
                        resp_new = requests.get(detail_url, headers={'headers':self.ua.random}, verify=False)
                        soup_new = BeautifulSoup(resp_new.content, 'html.parser')

                        date_ori = soup_new.find('div', {'class': 'flash-bg'}).find('div', {'class': 'info'}).find('time').find('span').get_text().strip()
                        # print(date_ori)
                        date = date_ori + ':00'
                        print(date)
                        date = datetime.datetime.strptime(str(date).split('.')[0], "%Y-%m-%d %H:%M:%S")

                        if self.args.start <= date <= self.args.end:
                            clean_md, clean_text, clean_doc, redu, reci, title, date, url, dianzan, \
                                pinglun, yuedu = self.html2res2(detail_url, date, str(idd))
                            if self.args.keyword == '' or self.args.keyword in clean_text:
                                save_file(self.args, idd, clean_md, clean_text, clean_doc, redu, reci, title, date, url, dianzan, pinglun, yuedu)
                                idd += 1
                            else:
                                print('该文章不包含关键词：{}！跳过！'.format(self.args.keyword))
                        elif date > self.args.end:
                            print('当前快讯比查找范围最新日期还要新，继续查找')
                            continue
                        else:
                            print('当前快讯过早，结束查找！')
                            flag = True
                            break
                    except Exception as e:
                        print(e)
                        print(detail_url + '  download failure !!!')
                        continue
                print(f"第{i}页爬取成功\n")
                i += 1
        except Exception as e:
            print('整体快讯网页爬取异常：%s' % e)
            print(traceback.format_exc())
            pass

        if os.path.exists("./tmp/img/"):
            shutil.rmtree("./tmp/img/")
        os.makedirs("./tmp/img/", exist_ok = True) 
        return
