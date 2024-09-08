import os
import time
import requests
import re
import urllib
import urllib.request
from urllib import request
import base64
import yaml
import traceback
import random
import os  
import threading
import json
import copy
import datetime
import magic
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from docx import Document 
import io
from pathlib import Path
from PIL import Image
from html2text import HTML2Text
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT 
from docx.oxml.shared import OxmlElement 
from docx.oxml.ns import qn 
from fake_useragent import UserAgent
from bs4 import BeautifulSoup
# noinspection PyUnresolvedReferences
from urllib.parse import urlparse
# from logger import logger


allinfo = {
        '新能源汽车':['1766033643315834880','SEI'],
        '人工智能':['1766033379397664768','SEI'],
        '集成电路':['1766034156832903168','SEI'],
        '工业母机':['1766033463598276608','SEI'],
        '新能源':['1766033744289464320','SEI'],
        '新一代移动通信':['1766034093557596160','SEI'],
        '生物技术':['1766033589087674368','SEI'],
        '工业软件':['1766034205675507712','SEI'],
        '新材料':['1766033523736182784','SEI'],

        '未来能源':['1766037300157493248','FI'],
        '未来健康':['1766037619054665728','FI'],
        '未来信息':['1766034292099186688','FI'],
        '未来网络':['1766037094787706880','FI'],
        '未来空间':['1766037771639222272','FI'],
        '未来制造':['1766037449856417792','FI'],

        '商业航天':['1787741534431309824','OI'],
        '低空经济':['1787751043652476928','OI']
    }


#text文本存入文件
def savetext2file(text, source, aid, day):
    path = 'filecontentic/{}/{}'.format(day, source)
    os.makedirs(path, exist_ok = True) 
    file = os.path.join(path, str(aid)+'.txt')
    with open(file, 'w', encoding='utf-8') as fout:
        fout.write(text + '\n')

def download(file_path, picture_url):
    # ua = UserAgent()
    # headers={'headers':ua.random}
    
    # r = requests.get(picture_url, headers=headers).content

    # with open(file_path, 'wb') as f:
    #     f.write(r)

    urllib.request.urlretrieve(picture_url,file_path)
 
#修改缩略图，300*200px
def resize_image(input_image_path, output_image_path, size=(300, 200)):
    with Image.open(input_image_path) as image:
        image.thumbnail(size)
        image.save(output_image_path)
        
def resize_image_1000(input_image_path, output_image_path):
    with Image.open(input_image_path) as image:
        width = image.size[0]
        height = image.size[1]
        scale = 1000/width
        n_img = image.resize((int(width*scale), int(height*scale)), Image.ANTIALIAS)
        save_image(n_img, output_image_path)

def image_to_base64(image_path):
    with open(image_path, "rb") as image_file:
        image_data = image_file.read()
        base64_data = base64.b64encode(image_data)
        base64_string = base64_data.decode("utf-8")
        return base64_string

def save_base64(image_path, base64path):
    image_base64 = image_to_base64(image_path)
    text = '''data:image/{};base64,{}'''.format(image_path.split('.')[-1],image_base64)
    with open(base64path,'w') as f:
        f.write(text)

def resize_image_rb(rb, size):
    '''修改二进制流图像，300*'''
    imagePixmap = rb.size
    n_img = rb.resize(size, Image.ANTIALIAS)
    return n_img


def save_image(img, save_path):
    try:
        img.save(save_path)
    except:
        img=img.convert('RGB')
        img.save(save_path)
    return

def thumbnail(image_path, resize_image_path, result_path):
    '''修改缩略图，并将base64值存入txt文件'''
    image_stream = Image.open(image_path)
    image_stream = image_stream.convert('P')
    thub_imgData = resize_image_rb(image_stream, (300,200))
    save_image(thub_imgData, resize_image_path)
    save_base64(resize_image_path, result_path)

def get_pages_content(url):
    try:
        # start = time.time()
        # service = Service(executable_path='/usr/local/bin/chromedriver-linux64/chromedriver')
        service = Service(executable_path=r'C:\Program Files\Google\Chrome\Application\chromedriver.exe')
        options = webdriver.ChromeOptions()
        options.add_argument('--headless')
        options.add_argument('--no-sandbox')  # fix:DevToolsActivePort file doesn't exist
        options.add_argument('--disable-gpu')  # fix:DevToolsActivePort file doesn't exist
        options.add_argument('--disable-dev-shm-usage')  # fix:DevToolsActivePort file doesn't exist
        options.add_argument('--remote-debugging-port=9222')  # fix:DevToolsActivePort file doesn't
        options.ignore_local_proxy_environment_variables()
        browser = webdriver.Chrome(service=service, options=options)
        # browser = webdriver.Chrome(options=options)
        browser.get(url)
        # print('1',time.time()-start)
        time.sleep(random.random() + 3)
        # print('2',time.time()-start)
        page_context = copy.deepcopy(browser.page_source)
        # print('3',time.time()-start)
        browser.close()
        # print('4',time.time()-start)
        return page_context
    except Exception as e:
        print(e)
        return None
    
    
#替换标题英文逗号
def rep_comma(text):
    text = str(text).strip('\n').strip().replace('\n','')
    return re.sub(',', '，', text)

def replace_symbol(text):
    '''
    处理文本中特殊符号，可以自定义
    '''
    pattern = '(?:[−⇔㎲・])'
    text = str(text).strip().replace('\n','').replace('*','').replace('\- ','- ').\
    replace('● ','- ')
    text = re.sub(pattern,'',text)
    return text


def remove_duplicates(lst):
    new_lst = [lst[0]]
    for i in range(1, len(lst)):
        if lst[i] != new_lst[-1]:
            new_lst.append(lst[i])
    return new_lst

def countdown(seconds, func):
    """
    定时器，用于检查函数是否在指定秒数内返回。
    """
    def decorator(*args, **kwargs):
        timer = threading.Timer(seconds, lambda: func(*args, **kwargs))
        try:
            timer.start()
            result = func(*args, **kwargs)
            return result
        except Exception as e:
            # 如果函数在3秒内引发异常，重新抛出异常
            timer.cancel()
            raise e
        finally:
            timer.cancel()
    return decorator

def ic_keywords():
    ic_keywords = {
                    '央企': ['中国电子', '中国电科', '中国信科', '华润集团',
                           '中国石化', '中核集团', '兵器工业', '国家电网',
                           '航空工业', '中国中车', '有研集团', '中国船舶',
                           '中国节能', '中国中化', '中国五矿'],
                    '国内外行业头部企业': ['台积电', '英伟达', '阿斯麦尔', '德州仪器', '高通', '英特尔', '英伟达',
                                          '超威半导体', '美光', '东京电子', '三星', 'SK海力士', '联发科',
                                          '新思科技', '楷登电子', '意法半导体', '英飞凌', 'ARM公司',
                                          '科磊', '泛林半导体', '日立', '华为', '平头哥半导体', '龙芯中科',
                                          '中芯国际', '华虹半导体', '北方华创', '上海微电子', '中微公司',
                                          '长电科技', '兆易创新', '北京君正']
                    }
    return ic_keywords


def image_to_jpg(image_path_or_stream):
    """使用Pillow将不支持的图片文件转换为.jpg文件"""
    f = io.BytesIO()
    if isinstance(image_path_or_stream, str):
        path = Path(image_path_or_stream)
        if path.suffix in {'.jpg', '.png', '.jfif', '.exif', '.gif', '.tiff', '.bmp'}:
            f = open(image_path_or_stream, mode='rb')
        else:
            Image.open(image_path_or_stream).convert('RGB').save(f, format='JPEG')
    else:
        buffer = image_path_or_stream.read()
        mime_type = magic.from_buffer(buffer, mime=True)
        if mime_type in {'image/jpeg', 'image/png', 'image/gif', 'image/tiff', 'image/x-ms-bmp'}:
            f = image_path_or_stream
        else:
            Image.open(io.BytesIO(buffer)).convert('RGB').save(f, format='JPEG')
    return f


def clean_html(args, idd, html, start_words=[], stop_words=[], element=['p','h2','h3','section','div'], img_links = ['data-src', 'src', 'data-original']):
    '''
    idd：编号
    html：主要的html内容
    start_words：获取到该词在内容中就开始处理
    stop_words：获取到该词在内容中就停止处理
    element：处理html时需要处理的所有元素
    img_links：处理图片时需要处理的所有元素
    '''
    del_imgsrc = ['https://mmbiz.qpic.cn/mmbiz_png/JWVrUoA6c7jJcRQELLRjftia53rn3icpaCLcIEbsDicWeU7cNNquUa9LKmic0Q1Kf6avCic5CRDImyj8C230H6Lb2AQ/640?wx_fmt=png',
                  'https://mmbiz.qpic.cn/mmbiz_png/JWVrUoA6c7gJJia5xWW774Yib26tmR3yMgHzNW2oAsP20CJe7EEWnXIWeBxsrpufUicc5H2QHFZ2ojjRIlrNzicrYA/640?wx_fmt=png']
    os.makedirs("./output/{}/{}/doc/{}/".format(args.day, args.industry, args.name), exist_ok = True)
    os.makedirs("./output/{}/{}/mdFiles/{}/".format(args.day, args.industry, args.name), exist_ok = True)
    os.makedirs("./output/{}/{}/text/{}/".format(args.day, args.industry, args.name), exist_ok = True)
    os.makedirs("./output/{}/{}/thumbFiles/{}/".format(args.day, args.industry, args.name), exist_ok = True) 
        
    res = []
    onlytext_res = []
    n = 1
    output_doc = Document() 
    
    # 设置默认字体、字号和中文字体
    font_name = u'宋体'
    output_doc.styles['Normal'].font.size = Pt(12)  # 设置默认字号为12号字体
    output_doc.styles['Normal'].font.name = font_name  # 设置默认字体为楷体
    output_doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # 设置中文字体为宋体

    paragraph = output_doc.add_paragraph()

    flag = False
    startflag = False
    content_set = set()
    for item in html.find_all(name=element):
        if flag:
            break
        if '<img' in str(item):
            try:           
                if args.name == '中国能源新闻网':
                    imgsrc = item.find('img')['src']
                    imgsrc = 'https://cpnn.com.cn/news/xny/'+'{}/'.format(imgsrc[4:10])+imgsrc[2:]
                    pre_img_type = imgsrc.split('.')[-1]
                elif args.name == '搜狐网':
                    imgsrc = item.find('img')['src']
                    imgsrc = 'https:'+imgsrc
                    pre_img_type = imgsrc.split('.')[-1]
                elif args.name == '中国储能网':
                    imgsrc = item.find('img')['src']
                    if 'https://www.escn.com.cn/' not in imgsrc:
                        imgsrc = 'https://www.escn.com.cn/'+imgsrc[:8]+'/'+imgsrc.split('_')[0][8:]+'/'+imgsrc
                    pre_img_type = imgsrc.split('.')[-1]
                elif args.name == '汽车之家':
                    imgsrc = item.find('img')['data-src']
                    imgsrc = 'https:'+imgsrc
                    pre_img_type = imgsrc.split('.')[-1]
                elif args.name == '电子工程专辑':
                    imgsrc = item.find('img')['src']
                    if 'https://www.eet-china.com' not in imgsrc:
                        imgsrc = 'https://www.eet-china.com'+imgsrc
                    pre_img_type = imgsrc.split('.')[-1]
                elif args.name == '新华网':
                    imgsrc = item.find('img')['src']
                    imgsrc = 'http://www.news.cn/tech/'+imgsrc[:8]+'/'+imgsrc[8:40]+'/'+imgsrc
                    pre_img_type = imgsrc.split('.')[-1]
                elif args.name == '中国核电信息网':
                    imgsrc = item.find('img')['src']
                    imgsrc = 'http://www.heneng.net.cn' + imgsrc
                    pre_img_type = imgsrc.split('.')[-1]
                elif args.name == '中国氢能源网':
                    imgsrc = item.find('img')['src']                  
                    imgsrc = 'http://www.china-hydrogen.org' + imgsrc
                    pre_img_type = imgsrc.split('.')[-1]
                elif args.name == '国家自然科学基金委员会':
                    imgsrc = item.find('img')['src']
                    imgsrc = 'https://www.nsfc.gov.cn'+imgsrc
                    pre_img_type = imgsrc.split('.')[-1]
                elif args.name == 'chinaev':
                    imgsrc = item.find('img')['src']
                    imgsrc = 'http://www.chinanev.net/' + imgsrc
                    pre_img_type = imgsrc.split('.')[-1]
                elif args.name == 'huawei':
                    imgsrc = item.find('img')['src']
                    imgsrc = 'https://digitalpower.huawei.com/' + imgsrc
                    pre_img_type = imgsrc.split('.')[-1]
                else:
                    reg = '(http[s]?://[0-9a-zA-Z_\-=?,/\.:~%&]*)'
                    imgreg = re.compile(reg)
                    imgsrc = ''
                    imgflag = False
                    for linksign in img_links:
                        if imgflag:
                            break
                        try:
                            imgsrc = item.find('img')[linksign]
                            if imgreg.search(imgsrc) is not None:
                                imgflag = True
                                break
                            else:
                                imgsrc = ''
                        except:
                            pass
                    # print(imgsrc)
                    if imgsrc in del_imgsrc:
                        continue
                    pattern = '(?:\.(?:jpg|png|jpeg|webp|jiff|JPG))'
                    if re.compile(pattern).search(imgsrc) is not None:
                        pre_img_type = re.compile(pattern).search(imgsrc).group().strip('.')
                    else:
                        pre_img_type = 'png'

                if imgsrc == '' or imgsrc in content_set or pre_img_type in ['gif']:
                    print('获取不到图片，或图片格式不在处理范围！')
                    continue
                
                # 如果图片以webp结尾，则将后缀去掉
                if imgsrc.endswith('.webp'):
                    imgsrc = imgsrc.split('.webp')[0]
                content_set.add(imgsrc)

                print('下载第{}张图片:{}'.format(n, imgsrc))
                img_type = pre_img_type if pre_img_type != '' else 'png'
                file_path = './tmp/img/{}_pre_{}.{}.{}'.format(args.name,idd,n,img_type)
                resize_file_path = './tmp/img/{}_{}.{}.{}'.format(args.name,idd,n,img_type)


                download(file_path, imgsrc)
                resize_image_1000(file_path, resize_file_path)
                resize_base64_str = image_to_base64(resize_file_path)
                imagetext = '![image](data:image/{};base64,{})'
                image_base64text = imagetext.format(resize_file_path.split('.')[-1],resize_base64_str)

                res.append(image_base64text)
                output_doc.add_picture(resize_file_path)
                paragraph = output_doc.add_paragraph()

                # 把第一张存为缩略图
                if n == 1:
                    image_path = file_path
                    if args.thub_type == 'txt':
                        resize_image_path = './tmp/img/{}_{}.{}'.format(args.name,idd,img_type)
                        result_path = './output/{}/{}/thumbFiles/{}/{}.txt'.format(args.day,
                                                                                args.industry,
                                                                                args.name,idd)
                        thumbnail(image_path, resize_image_path, result_path)
                    else:
                        resize_image_path = './output/{}/{}/thumbFiles/{}/{}.{}'.format(args.day,
                                                                                        args.industry,
                                                                                        args.name,idd,img_type)
                        result_path = resize_image_path = './tmp/img/{}_{}.txt'.format(args.name,idd)
                        thumbnail(image_path, resize_image_path, result_path)
                n += 1
            except Exception as e:
                print('处理图片出现错误：{}'.format(e))
                pass
        elif 'video' in str(item) or '.gif' in str(item):
            continue
        
        #　处理文字内容
        try:
            content = item.get_text().strip('\n').strip()
            content = replace_symbol(content)

            # 如遇网址则跳过，可注释
            httpreg = '(http[s]?://[0-9a-zA-Z_\-=?,/\.:~%&]*)'
            if re.compile(httpreg).search(content) is not None:
                continue

            if content in content_set:
                continue
            content_set.add(content)

            if len(start_words) == 0:
                startflag = True
            else:
                for start_word in start_words:
                    if start_word in content:
                        startflag = True

            if startflag:
                for stop_word in stop_words:
                    if stop_word in content:
                        flag = True
                        break
                
                if flag:
                    break
                onlytext_res.append(content)
                if str(item).startswith('<h2') or str(item).startswith('<h3'):
                    res.append('## <font color="#00e4ff">**{}**</font>'.format(content))

                    r = paragraph.add_run(content)
                    r.font.name = font_name
                    r.bold = True

                    r = paragraph.add_run('\n\n\n')
                    r.font.name = font_name
                else:
                    res.append(content)

                    r = paragraph.add_run(content+'\n\n\n')
                    r.font.name = font_name
        except Exception as e:
            print('处理文字出现错误：{}'.format(e))
            pass
    res_text = '\n\n'.join(res)
    onlytext_res_text = '\n\n'.join(onlytext_res)
    return res_text, onlytext_res_text, output_doc


def clean_markdown(args, idd, md, start_words=[], stop_words=[]):
    '''清洗markdown源数据，下载图片等
    idd：编号
    md：需要清洗的markdown
    start_words：获取到该词在内容中就开始处理
    stop_words：获取到该词在内容中就停止处理
    '''

    os.makedirs("./output/{}/{}/doc/{}/".format(args.day, args.industry, args.name), exist_ok = True)
    os.makedirs("./output/{}/{}/mdFiles/{}/".format(args.day, args.industry, args.name), exist_ok = True)
    os.makedirs("./output/{}/{}/text/{}/".format(args.day, args.industry, args.name), exist_ok = True)
    os.makedirs("./output/{}/{}/thumbFiles/{}/".format(args.day, args.industry, args.name), exist_ok = True) 

    text_list = [text.strip('\n').strip() for text in md.split('\n') if text.strip('\n').strip() != '']
    res = []
    onlytext_res = []
    n = 1
    output_doc = Document() 
    paragraph = output_doc.add_paragraph()

    # 设置默认字体、字号和中文字体
    font_name = u'宋体'
    output_doc.styles['Normal'].font.size = Pt(12)  # 设置默认字号为12号字体
    output_doc.styles['Normal'].font.name = font_name  # 设置默认字体为楷体
    output_doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # 设置中文字体为宋体
    
    flag = False
    startflag = False
    content_set = set()
    for text in text_list:
        pattern = '(?:jpg|png|jpeg|webp|jiff|JPG)'
        if re.compile(pattern).search(text) is not None:                

            try:
                if '中国新材料产业技术创新平台' in args.name:
                    if ('https' in text) or ('http' in text):
                        reg = '(http(s)?://[0-9a-zA-Z_/\%.…-]*?\.jpg)|(http(s)?://[0-9a-zA-Z_/\%.…-]*?\.png)|(http(s)?://[0-9a-zA-Z_/\%.…-]*?\.jpeg)'
                        imgreg = re.compile(reg)
                        if len(re.findall(imgreg, text)) > 0:
                            picture_turple = re.findall(imgreg, text)[0]
                            picture_url = [x.strip() for x in list(picture_turple) if x.strip() != ''][0]
                        else:
                            picture_url = ''
                    else:
                        reg = "([0-9a-zA-Z_/\.-]*?\.jpg)|([0-9a-zA-Z_/\.-]*?\.jpeg)|([0-9a-zA-Z_/\.-]*?\.png)"
                        imgreg = re.compile(reg)
                        picture_turple = re.findall(imgreg, text)[0]
                        pic_url = [x.strip() for x in list(picture_turple) if x.strip() != ''][0]
                        picture_url = 'http://www.chinanmia.com' + pic_url  
                else:
                    reg = '(http(s)?://[0-9a-zA-Z_/\.-]*?\.(?:jpg|png|jpeg|webp|jiff|JPG))'
                    imgreg = re.compile(reg)
                    if len(re.findall(imgreg, text)) > 0:
                        picture_turple = re.findall(imgreg, text)[0]
                        picture_url = [x.strip() for x in list(picture_turple) if x.strip() != ''][0]
                    else:
                        picture_url = ''

                if picture_url == '' or picture_url in content_set:
                    print('获取不到图片，或图片格式不在处理范围！')
                    continue
                
                content_set.add(picture_url)

                print('下载第{}张图片:{}'.format(n, picture_url))
                try:
                    img_type = picture_url.split('.')[-1]
                    if re.compile(pattern).search(img_type) is not None:
                        img_type = img_type
                    else:
                        img_type = 'png'
                except:
                    img_type = 'png'
                file_path = './tmp/img/{}_pre_{}.{}.{}'.format(args.name,idd,n,img_type)
                resize_file_path = './tmp/img/{}_{}.{}.{}'.format(args.name,idd,n,img_type)
                download(file_path, picture_url)
                resize_image_1000(file_path, resize_file_path)
                resize_base64_str = image_to_base64(resize_file_path)
                imagetext = '''![image](data:image/{};base64,{})'''
                image_base64text = imagetext.format(resize_file_path.split('.')[-1],resize_base64_str)
                
                res.append(image_base64text)
                output_doc.add_picture(resize_file_path)
                paragraph = output_doc.add_paragraph()

                # 把第一张存为缩略图
                if n == 1:
                    image_path = file_path
                    if args.thub_type == 'txt':
                        resize_image_path = './tmp/img/{}_thub_{}.{}.png'.format(args.name,idd,n)

                        result_path = './output/{}/{}/thumbFiles/{}/{}.txt'.format(args.day,
                                                                                    args.industry,
                                                                                    args.name,idd)
                        thumbnail(image_path, resize_image_path, result_path)
                    else:
                        resize_image_path = './output/{}/{}/thumbFiles/{}/{}.png'.format(args.day,
                                                                                        args.industry,
                                                                                        args.name,idd)
                        result_path = resize_image_path = './tmp/img/{}_thub_{}.{}.txt'.format(args.name,idd,n)

                        thumbnail(image_path, resize_image_path, result_path)
                n += 1

            except Exception as e:
                print('处理图片出现错误：{}'.format(e))
                print(traceback.format_exc())
                pass
        
        elif '.gif' in text:
            continue

        # 处理文字信息
        try:
            content = replace_symbol(text)

            # 如遇网址则跳过，可注释
            httpreg = '(http[s]?://[0-9a-zA-Z_\-=?,/\.:~%&]*)'
            if re.compile(httpreg).search(content) is not None:
                continue

            if content in content_set:
                continue
            content_set.add(content)


            if len(start_words) == 0:
                        startflag = True
            else:
                for start_word in start_words:
                    if start_word in content:
                        startflag = True

            if startflag:
                for stop_word in stop_words:
                    if stop_word in content:
                        flag = True
                        break
                
                if flag:
                    break
                onlytext_res.append(content)
                if str(content).startswith('## '):
                    res.append('## <font color="#00e4ff">**{}**</font>'.format(content))
                    
                    r = paragraph.add_run(content)
                    r.font.name = font_name
                    r.bold = True

                    r = paragraph.add_run('\n\n\n')
                    r.font.name = font_name
                elif str(content).startswith('### '):
                    res.append('### <font color="#00e4ff">**{}**</font>'.format(content))
                    
                    r = paragraph.add_run(content)
                    r.font.name = font_name
                    r.bold = True
                    
                    r = paragraph.add_run('\n\n\n')
                    r.font.name = font_name
                else:
                    res.append(content)

                    r = paragraph.add_run(content+'\n\n\n')
                    r.font.name = font_name
        except Exception as e:
            print('处理文字出现错误：{}'.format(e))
            pass

    res_text = '\n\n'.join(res)
    onlytext_res_text = '\n\n'.join(onlytext_res)
    return res_text, onlytext_res_text, output_doc


def save_file(args, idd, clean_md, clean_text, clean_doc, redu, reci, title, date, url, dianzan, pinglun, yuedu, saveflag=True):
    '''存储爬取信息
    idd：编号
    clean_md：清洗后的markdown
    clean_text：清洗后的文本
    clean_doc：清洗后的doc
    redu：热度
    reci：热词
    title：标题
    date：时间
    url：网址
    dianzan：点赞量
    pinglun：评论量
    yuedu：阅读量
    saveflag：是否保存点赞量、评论量、阅读量信息（推荐时需要用，但仅存储单个网址信息时不用）
    '''

    os.makedirs("./output/{}/{}/doc/{}/".format(args.day, args.industry, args.name), exist_ok = True)
    os.makedirs("./output/{}/{}/mdFiles/{}/".format(args.day, args.industry, args.name), exist_ok = True)
    os.makedirs("./output/{}/{}/text/{}/".format(args.day, args.industry, args.name), exist_ok = True)
    os.makedirs("./output/{}/{}/thumbFiles/{}/".format(args.day, args.industry, args.name), exist_ok = True)    

    # 存放markdown
    markdownpath = "./output/{}/{}/mdFiles/{}/{}.txt".format(args.day, 
                                                            args.industry,
                                                            args.name, 
                                                            str(idd))
        
    # 存放text
    textpath = "./output/{}/{}/text/{}/{}.txt".format(args.day, 
                                                    args.industry,
                                                    args.name, 
                                                    str(idd))

    # 存放doc
    docpath = "./output/{}/{}/doc/{}/{}.docx".format(args.day, 
                                                    args.industry,
                                                    args.name, 
                                                    str(idd))

    with open(markdownpath, 'w', encoding='utf8') as my_file:
        my_file.write(clean_md)

    with open(textpath, 'w', encoding='utf8') as my_file:
        my_file.write(clean_text)

    clean_doc.save(docpath)

    if args.thub_type == 'txt':
        hasimg = 1 if os.path.exists("./output/{}/{}/thumbFiles/{}/{}.txt".format(
                args.day, args.industry, args.name, str(idd))) else 0
        thumbpath = "./output/{}/{}/thumbFiles/{}/{}.txt".format(args.day, args.industry, args.name, str(idd))
    else:
        hasimg = 1 if os.path.exists("./output/{}/{}/thumbFiles/{}/{}.png".format(
                args.day, args.industry, args.name, str(idd))) else 0
        thumbpath = "./output/{}/{}/thumbFiles/{}/{}.txt".format(args.day, args.industry, args.name, str(idd))
    
    
    if saveflag:
        
        # 修改日期为当天范围
        today = datetime.datetime.now().strftime('%Y-%m-%d')
        today_start_time = datetime.datetime.strptime(today + ' 00:00:00', "%Y-%m-%d %H:%M:%S")
        today_end_time = datetime.datetime.strptime(today + ' 12:00:00', "%Y-%m-%d %H:%M:%S")

        if isinstance(date, str):
            date = datetime.datetime.strptime(date, "%Y-%m-%d %H:%M:%S")

        if not today_start_time <= date <= today_end_time:
            date = today + ' 08:00:00'

        with open(args.info_path, "a+", encoding='utf8') as f:
            # f.write(f"{idd},{args.industry_name},{args.industry_type},"\
            #         f"{args.industry_ID}, {args.fuzeren} ,{args.dom_or_int}, {redu}, {reci} ,{title},"\
            #         f"{date},{args.source_name}, {url}, {dianzan}, {pinglun}, {yuedu}, {hasimg}\n")

            f.write(f"{idd},{args.industry_name},{args.industry_type},"\
                    f"{args.industry_ID}, {args.fuzeren} ,{args.dom_or_int}, {redu}, {reci} ,{title},"\
                    f"{date},{args.source_name}, {url}\n")
    

    print('doc已保存至{}!'.format(docpath))
    print('Markdown已保存至{}!'.format(markdownpath))
    print('文本已保存至{}!'.format(textpath))
    if hasimg:
        print('缩略图已保存至{}!'.format(thumbpath))
    print('信息已保存至{}!'.format(args.info_path))
    return 